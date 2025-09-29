from selenium import webdriver
from selenium.webdriver.common.by import By
from time import sleep
from dotenv import load_dotenv
import os
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.window import WindowTypes
from selenium.webdriver.chrome.options import Options
load_dotenv()
import pandas as pd
from PyPDF2 import PdfReader
import docx
import glob
import sys
import contextlib
import logging
import warnings
import shutil
import tempfile
import subprocess
try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    def rtf_to_text(rtf_str):
        return rtf_str

# silence noisy loggers and warnings from PDF/third-party libs
logging.getLogger("PyPDF2").setLevel(logging.CRITICAL)
logging.getLogger("pikepdf").setLevel(logging.CRITICAL)
warnings.filterwarnings("ignore")

@contextlib.contextmanager
def suppress_stderr():
    """Temporarily redirect sys.stderr to /dev/null to hide library noise."""
    with open(os.devnull, "w") as devnull:
        old_stderr = sys.stderr
        sys.stderr = devnull
        try:
            yield
        finally:
            sys.stderr = old_stderr

def init_driver(url: str) -> webdriver:
    download_dir = os.path.join(os.getcwd(), "selenium_downloads")
    os.makedirs(download_dir, exist_ok=True)
    chrome_options = Options()
    prefs = {"download.default_directory": download_dir,
             "download.prompt_for_download": False} 
    chrome_options.add_experimental_option("prefs", prefs)
    chrome_options.add_argument('--headless=new')
    driver = webdriver.Chrome(options=chrome_options)
    driver.get(url)
    return driver
# user must sign into terplink
def user_sign_in() -> None:
    shadow_host = driver.find_element(By.ID, "parent-root")
    shadow_root = driver.execute_script("return arguments[0].shadowRoot", shadow_host)
    header = driver.execute_script("return arguments[0].querySelector('header')", shadow_root)
    sign_in_button = driver.execute_script(
        "return arguments[0].querySelector('a.MuiButton-containedPrimary')", header
    )
    sign_in_button.click()

# loads all organizations
def load():
    while True: 
        sleep(0.5) 
        try:
            load_button = driver.find_element(By.XPATH, "//div[@class='outlinedButton']/button")
            load_button.click()

        except Exception:
            break
def collect_org_info() -> list:

    name = ""
    desc = ""
    additional_info = {
        "mission_statement": "", 
        "membership_requirements": "",
        "how_to_get_involved": "",
        "general_meeting_information": "",
        "expected_time_commitment": "",
        "meeting_schedule": ""
    }
    try: 
        name = driver.find_element(By.TAG_NAME, "h1").text
    except Exception:
        name = ""
    
    try: 
        desc = driver.find_element(By.CLASS_NAME, "bodyText-large.userSupplied").text
    except Exception:
        desc = ""
    
    # collects mission statement and cleans it
    try:

        additional_info["mission_statement"] =  ((driver.find_element(By.XPATH, "//strong[contains(text(), 'Mission Statement')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text
        additional_info["mission_statement"] = additional_info["mission_statement"].replace("Mission Statement\nAn organization's mission statement should provide a description of the group's purpose.\n", "")
        
        if additional_info["mission_statement"].strip() == "No Response":
            additional_info["mission_statement"] = None
    except Exception:
        additional_info["mission_statement"] = None
    

    # collects membership requirements and cleans it
    try: 
        additional_info["membership_requirements"] = ((driver.find_element(By.XPATH, "//strong[contains(text(), 'Membership Requirements')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text
        additional_info["membership_requirements"] = additional_info["membership_requirements"].replace("Membership Requirements\nPlease list information about the organizations selection process and include membership requirements if applicable.\n", "")
        if additional_info["membership_requirements"].strip() == "No Response":
            additional_info["membership_requirements"] = None
    except Exception:
        additional_info["membership_requirements"] = None


    # collects how to get involved and cleans it
    try:
        additional_info["how_to_get_involved"] = ((driver.find_element(By.XPATH, "//strong[contains(text(), 'How would a student get involved in your organization?')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text
        additional_info["how_to_get_involved"] = additional_info["how_to_get_involved"].replace('How would a student get involved in your organization?\nIncluding, but not limited to, "show up to our meetings!" or "you must apply and applications are available," or "email us and let us know you want to be involved!"\n', "")
        if additional_info["how_to_get_involved"].strip() == "No Response":
            additional_info["how_to_get_involved"] = None
    except Exception:
        additional_info["how_to_get_involved"] = None


    # collects general meeting information and when general meetings are and cleans it 
    try:
        additional_info["general_meeting_information"] = ((driver.find_element(By.XPATH, "//strong[contains(text(), 'Does your organization have general body meetings that are open for potential new members to attend?')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text
        additional_info["general_meeting_information"] = additional_info["general_meeting_information"].replace("Does your organization have general body meetings that are open for potential new members to attend?\n", "")
        if additional_info["general_meeting_information"].strip() == "No Response":
            additional_info["general_meeting_information"] = ""

    except Exception:
        additional_info["general_meeting_information"] = ""
    try:
        temp = ((driver.find_element(By.XPATH, "//strong[contains(text(), 'How often does your organization have general body meetings?')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text.strip()
        temp = temp.replace("How often does your organization have general body meetings?\n", "")
        if temp != "No Response":
            additional_info["general_meeting_information"] += temp
        
    except Exception:
        additional_info["general_meeting_information"] += ""
    
    if additional_info["general_meeting_information"] == "":
        additional_info["general_meeting_information"] = None


    # collects expected time commitment and cleans it 
    try: 
        additional_info["expected_time_commitment"] = ((driver.find_element(By.XPATH, "//strong[contains(text(), 'How would you rate the time commitment level to be a member of your organization?')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text
        additional_info["expected_time_commitment"] = additional_info["expected_time_commitment"].replace("How would you rate the time commitment level to be a member of your organization? \n", "")
        if additional_info["expected_time_commitment"].strip() == "No Response":
            additional_info["expected_time_commitment"] = None
    except Exception:
        additional_info["expected_time_commitment"] = None

    # collects meeting schedule and cleans it
    try:
        additional_info["meeting_schedule"] = ((driver.find_element(By.XPATH, "//strong[contains(text(), 'When are most of the commitments')]")).find_element(By.XPATH, "./parent::div").find_element(By.XPATH, "./parent::div")).text
        additional_info["meeting_schedule"] = additional_info["meeting_schedule"].replace("When are most of the commitments (meetings, events, office hours, etc.) for your organization? (check all that apply)\n", "")
        if additional_info["meeting_schedule"].strip() == "No Response":
            additional_info["meeting_schedule"] = None
    except Exception:
        additional_info["expected_time_commitment"] = None
    return [name, desc, additional_info]

def get_news_info(url) -> list:
    articles = []
    news_url = url + '/news'
    original_window = driver.current_window_handle
    driver.switch_to.new_window(WindowTypes.TAB)
    driver.get(news_url)
    load()
    news = driver.find_elements(By.XPATH, "//div[@id='news-list']//a")
    if len(news) == 0:
        driver.close()
        driver.switch_to.window(original_window)
        return []
    for i in news:
        href = i.get_attribute("href")
        news_window = driver.current_window_handle
        driver.switch_to.new_window(WindowTypes.TAB)
        driver.get(href)
        

        try:
            article = (driver.find_element(By.CLASS_NAME, "MuiCardContent-root.articleText")).text
            articles.append(article)
        except Exception:
            print("no article")
        driver.close()
        driver.switch_to.window(news_window)
    driver.close()
    driver.switch_to.window(original_window)
    return articles

def get_events_info(url) -> list:
    event_info = []
    event_url = url + '/events'
    original_window = driver.current_window_handle
    driver.switch_to.new_window(WindowTypes.TAB)
    driver.get(event_url)
    load()
    events = driver.find_elements(By.XPATH, "//div[@id='org-event-discovery-list']//a")
    if len(events) == 0:
        driver.close()
        driver.switch_to.window(original_window)
        return []
    for event in events:
        href = event.get_attribute("href")
        events_window = driver.current_window_handle
        driver.switch_to.new_window(WindowTypes.TAB)
        driver.get(href)
        try:
            event_info.append(driver.find_element(By.XPATH, "//div[@class='DescriptionText']").text)
        except Exception:
            print("cannot find event")
        driver.close()
        driver.switch_to.window(events_window)

    driver.close()
    driver.switch_to.window(original_window)
    return event_info

def get_document():
    try:
        document_url = driver.find_element(By.CSS_SELECTOR, '[aria-label="download Constitution/Bylaws"]')
    except Exception:
        return ""
    document_url.click()
    sleep(2)
    file_path = None
    download_dir = os.path.join(os.getcwd(), "selenium_downloads")
    for i in range(30):
        files = glob.glob(os.path.join(download_dir, '*'))
        if files:
            file_path = max(files, key=os.path.getctime)

            if not file_path.endswith('.crdownload'):
                break
        sleep(0.5)
    
    if not file_path or file_path.endswith('.crdownload'):
        print('no')
        return ""
    
    text = ""
    try:
        if file_path.lower().endswith('.pdf'):
            with suppress_stderr():
                with open(file_path, "rb") as f:
                    reader = PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
        elif file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + " "
        elif file_path.lower().endswith('.doc'):
            # Try to convert .doc -> .docx using LibreOffice (soffice) if available
            soffice = shutil.which('soffice') or shutil.which('libreoffice')
            if soffice:
                with tempfile.TemporaryDirectory() as tmpdir:
                    try:
                        subprocess.run([soffice, '--headless', '--convert-to', 'docx', file_path, '--outdir', tmpdir], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        # find converted file
                        base = os.path.splitext(os.path.basename(file_path))[0]
                        converted = os.path.join(tmpdir, base + '.docx')
                        if os.path.exists(converted):
                            doc = docx.Document(converted)
                            for para in doc.paragraphs:
                                text += para.text + " "
                        else:
                            print(f"Conversion produced no output for {file_path}")
                            text = ""
                    except Exception as e:
                        print(f"Failed to convert .doc to .docx for {file_path}: {e}")
                        text = ""
            else:
                print(f"LibreOffice (soffice) not found; cannot convert .doc: {file_path}")
                text = ""
        elif file_path.lower().endswith('.rtf') or file_path.lower().endswith('.docx.rtf'):
            try:
                # open as binary and decode to text for rtf_to_text
                with open(file_path, 'rb') as f:
                    raw = f.read()
                try:
                    rtf_str = raw.decode('utf-8')
                except Exception:
                    try:
                        rtf_str = raw.decode('latin-1')
                    except Exception:
                        rtf_str = raw.decode(errors='ignore')
                text = rtf_to_text(rtf_str)
            except Exception as e:
                print(f'Could not extract .rtf file: {file_path}, error: {e}')
                text = ""
        else:
            print("Unsupported file type:", file_path)
            text = ""
    finally:
        # remove the downloaded file to avoid accumulating files
        try:
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
        except Exception:
            pass

    return text
    

    

def collect_orgs() -> pd.DataFrame:
    df = pd.DataFrame(columns=['Name', 'Description', 'Additional Information', 'Events', 'Document', 'News', 'URL']) 
    # creates a list containing all orgs
    orgs = driver.find_elements(By.CLASS_NAME, "MuiCard-root")
    links = []
    # gather link for each org
    for org in orgs:
        try:
            links.append(org.find_element(By.XPATH, "./parent::a").get_attribute("href"))
        except Exception:
            print(Exception)


    original_window = driver.current_window_handle
    count = 1
    for link in links:
        print(count)
        print(link)
        driver.switch_to.new_window(WindowTypes.TAB)
        driver.get(link)
        sleep(.5)
        row = collect_org_info()
        row.append(get_events_info(link))
        row.append(get_document())
        row.append(get_news_info(link))
        row.append(link)


        df.loc[len(df)] = row
        driver.close()
        driver.switch_to.window(original_window)
        sleep(.5)
        count += 1
        
        
    return df


    




if __name__ == '__main__':
    url = "https://terplink.umd.edu/organizations"
    test = 'https://terplink.umd.edu/organization/salvation-and-praise'
    driver = init_driver(url)
    # waits until load button is available
    WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.XPATH, "//div[@class='outlinedButton']/button")))
    load() # presses load more button until all orgs are loaded
    df = collect_orgs() # navigates through each orgs page and collects information


    for col in ['Name', 'Description', 'Document']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.replace('\n', ' ', regex=False).str.replace('\r', ' ', regex=False)

    df.to_csv('../data/org-data-unprocessed.csv')
    driver.quit()
